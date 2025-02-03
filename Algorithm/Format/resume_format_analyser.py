import fitz  # PyMuPDF for PDFs
import docx  # python-docx for Word documents
import logging
from typing import Dict, List, Optional, Set, Tuple
import re
from dataclasses import dataclass
from collections import defaultdict
import os
from enum import Enum
import yaml
import json
from datetime import datetime
import numpy as np
from PIL import Image
import io
import textwrap

class DocumentType(Enum):
    PDF = "application/pdf"
    DOCX = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    UNKNOWN = "unknown"

@dataclass
class FormattingIssue:
    """Data class to store formatting issues"""
    severity: str  # 'high', 'medium', 'low'
    category: str  # 'font', 'spacing', 'structure', etc.
    description: str
    suggestion: str
    location: Optional[str] = None

@dataclass
class SectionAnalysis:
    """Data class to store section analysis results"""
    name: str
    content: str
    font_info: Dict[str, int]
    spacing: float
    bullet_points: int
    formatting_issues: List[FormattingIssue]

class ResumeFormatAnalyzer:
    def __init__(self, config_path: str = 'format_config.yaml'):
        """
        Initialize the Resume Format Analyzer with enhanced capabilities.
        
        Args:
            config_path (str): Path to configuration file
        """
        self.setup_logging()
        self.config = self.load_config(config_path)
        self.initialize_patterns()

    def setup_logging(self):
        """Configure logging system"""
        logging.basicConfig(
            level=logging.INFO,
            format='%(asctime)s - %(levelname)s - [%(filename)s:%(lineno)d] - %(message)s',
            handlers=[
                logging.FileHandler('Logs/resume_format_analyzer.log'),
                logging.StreamHandler()
            ]
        )
        self.logger = logging.getLogger(__name__)

    def load_config(self, config_path: str) -> Dict:
        """Load configuration from YAML file"""
        try:
            if os.path.exists(config_path):
                with open(config_path, 'r') as f:
                    return yaml.safe_load(f)
            return self.get_default_config()
        except Exception as e:
            self.logger.error(f"Error loading config: {str(e)}")
            return self.get_default_config()

    def get_default_config(self) -> Dict:
        """Provide default configuration"""
        return {
            'formatting_rules': {
                'max_fonts': 3,
                'min_margin': 0.5,  # inches
                'max_margin': 1.0,  # inches
                'recommended_line_spacing': 1.15,
                'max_bullet_indent': 0.5,  # inches
                'section_spacing': 1.0,  # inches
                'max_page_count': 2
            },
            'section_requirements': {
                'required': ['summary', 'experience', 'education', 'skills'],
                'optional': ['certifications', 'achievements', 'projects']
            },
            'style_guidelines': {
                'recommended_fonts': ['Arial', 'Calibri', 'Times New Roman', 'Helvetica'],
                'min_font_size': 10,
                'max_font_size': 14,
                'header_font_size': 16,
                'bullet_indent_levels': 2
            }
        }

    def initialize_patterns(self):
        """Initialize regex patterns and formatting rules"""
        self.bullet_patterns = {
            'standard': [r"•", r"▪", r"◦", r"-", r"—", r"→", r"✔", r"✓"],
            'numbered': [r"^\d+\.", r"^[a-z]\.", r"^\[\d+\]"],
            'special': [r"★", r"✦", r"◆", r"❖"]
        }
        
        self.section_patterns = {
            'headers': [
                re.compile(rf"\b{keyword}\b", re.IGNORECASE)
                for keyword in self.config['section_requirements']['required'] +
                             self.config['section_requirements']['optional']
            ],
            'contact_info': [
                re.compile(r'\b[\w\.-]+@[\w\.-]+\.\w+\b'),  # email
                re.compile(r'\b\d{3}[-.]?\d{3}[-.]?\d{4}\b'),  # phone
                re.compile(r'linkedin\.com/\S+')  # LinkedIn
            ]
        }

    def detect_document_type(self, file_path: str) -> DocumentType:
        """Detect document type from file extension and content"""
        try:
            _, ext = os.path.splitext(file_path.lower())
            if ext == '.pdf':
                return DocumentType.PDF
            elif ext in ['.docx', '.doc']:
                return DocumentType.DOCX
            return DocumentType.UNKNOWN
        except Exception as e:
            self.logger.error(f"Error detecting document type: {str(e)}")
            return DocumentType.UNKNOWN

    def analyze_pdf_formatting(self, file_path: str) -> Dict:
        """
        Analyze PDF formatting with enhanced capabilities
        
        Args:
            file_path (str): Path to PDF file
            
        Returns:
            Dict: Comprehensive formatting analysis
        """
        try:
            doc = fitz.open(file_path)
            analysis = {
                'document_info': self.get_pdf_document_info(doc),
                'sections': [],
                'formatting_issues': [],
                'overall_statistics': defaultdict(int),
                'visual_elements': []
            }

            # Analyze each page
            for page_num, page in enumerate(doc):
                page_analysis = self.analyze_pdf_page(page, page_num)
                analysis['sections'].extend(page_analysis['sections'])
                analysis['formatting_issues'].extend(page_analysis['issues'])
                analysis['visual_elements'].extend(page_analysis['visual_elements'])
                
                # Update statistics
                for key, value in page_analysis['statistics'].items():
                    analysis['overall_statistics'][key] += value

            # Post-process analysis
            analysis['recommendations'] = self.generate_formatting_recommendations(analysis)
            analysis['compliance_score'] = self.calculate_compliance_score(analysis)
            
            return analysis

        except Exception as e:
            self.logger.error(f"Error analyzing PDF formatting: {str(e)}")
            return None

    def analyze_pdf_page(self, page, page_num: int) -> Dict:
        """Analyze individual PDF page"""
        analysis = {
            'sections': [],
            'issues': [],
            'statistics': defaultdict(int),
            'visual_elements': []
        }

        # Extract text blocks and analyze formatting
        blocks = page.get_text("dict")["blocks"]
        
        for block in blocks:
            if block.get("type") == 0:  # Text block
                section = self.analyze_text_block(block, page_num)
                if section:
                    analysis['sections'].append(section)
            elif block.get("type") == 1:  # Image block
                image_info = self.analyze_image_block(block, page)
                analysis['visual_elements'].append(image_info)

        # Analyze spacing and layout
        layout_issues = self.analyze_page_layout(page)
        analysis['issues'].extend(layout_issues)

        return analysis

    def analyze_text_block(self, block: Dict, page_num: int) -> Optional[SectionAnalysis]:
        """Analyze text block formatting"""
        try:
            text_content = []
            font_info = defaultdict(int)
            spacing_values = []
            bullet_count = 0

            for line in block.get("lines", []):
                for span in line.get("spans", []):
                    text_content.append(span["text"])
                    font_info[span["font"]] += 1
                    
                    # Check for bullets
                    if any(re.match(pattern, span["text"].strip()) 
                          for patterns in self.bullet_patterns.values() 
                          for pattern in patterns):
                        bullet_count += 1
                    
                    # Track line spacing
                    if "size" in span:
                        spacing_values.append(span["size"])

            # Detect section type
            content = " ".join(text_content)
            section_type = self.detect_section_type(content)
            
            if section_type:
                return SectionAnalysis(
                    name=section_type,
                    content=content,
                    font_info=dict(font_info),
                    spacing=np.mean(spacing_values) if spacing_values else 0,
                    bullet_points=bullet_count,
                    formatting_issues=self.check_section_formatting(content, font_info, spacing_values)
                )
            
            return None

        except Exception as e:
            self.logger.error(f"Error analyzing text block: {str(e)}")
            return None

    def analyze_docx_formatting(self, file_path: str) -> Dict:
        """
        Analyze DOCX formatting with enhanced capabilities
        
        Args:
            file_path (str): Path to DOCX file
            
        Returns:
            Dict: Comprehensive formatting analysis
        """
        try:
            doc = docx.Document(file_path)
            analysis = {
                'document_info': self.get_docx_document_info(doc),
                'sections': [],
                'formatting_issues': [],
                'overall_statistics': defaultdict(int),
                'style_consistency': {}
            }

            current_section = None
            section_content = []
            
            for para in doc.paragraphs:
                # Analyze paragraph formatting
                para_analysis = self.analyze_docx_paragraph(para)
                
                # Section detection
                if self.is_section_header(para.text):
                    if current_section:
                        section = self.create_section_analysis(
                            current_section, section_content)
                        analysis['sections'].append(section)
                    current_section = para.text
                    section_content = []
                else:
                    section_content.append(para_analysis)

                # Update statistics
                self.update_docx_statistics(analysis['overall_statistics'], para_analysis)

            # Add last section
            if current_section and section_content:
                section = self.create_section_analysis(current_section, section_content)
                analysis['sections'].append(section)

            # Post-process analysis
            analysis['style_consistency'] = self.analyze_style_consistency(doc)
            analysis['recommendations'] = self.generate_formatting_recommendations(analysis)
            analysis['compliance_score'] = self.calculate_compliance_score(analysis)

            return analysis

        except Exception as e:
            self.logger.error(f"Error analyzing DOCX formatting: {str(e)}")
            return None

    def generate_formatting_recommendations(self, analysis: Dict) -> List[str]:
        """Generate specific formatting recommendations"""
        recommendations = []
        
        # Font recommendations
        if len(analysis['overall_statistics'].get('unique_fonts', [])) > self.config['formatting_rules']['max_fonts']:
            recommendations.append(
                "Reduce the number of different fonts used. Stick to 2-3 professional fonts."
            )
        
        # Spacing recommendations
        if analysis['overall_statistics'].get('average_line_spacing', 0) < self.config['formatting_rules']['recommended_line_spacing']:
            recommendations.append(
                "Increase line spacing for better readability."
            )
        
        # Section structure recommendations
        missing_sections = set(self.config['section_requirements']['required']) - set(s.name for s in analysis.get('sections', []))
        if missing_sections:
            recommendations.append(
                f"Add missing required sections: {', '.join(missing_sections)}"
            )
        
        return recommendations

    def calculate_compliance_score(self, analysis: Dict) -> float:
        """Calculate overall formatting compliance score"""
        scores = {
            'font_consistency': self._calculate_font_score(analysis),
            'spacing_consistency': self._calculate_spacing_score(analysis),
            'section_structure': self._calculate_section_score(analysis),
            'visual_appeal': self._calculate_visual_score(analysis)
        }
        
        weights = {
            'font_consistency': 0.3,
            'spacing_consistency': 0.2,
            'section_structure': 0.3,
            'visual_appeal': 0.2
        }
        
        return sum(score * weights[category] for category, score in scores.items())

    def generate_report(self, analysis: Dict) -> Dict:
        """
        Generate detailed formatting analysis report
        
        Args:
            analysis (Dict): Analysis results
            
        Returns:
            Dict: Formatted report
        """
        return {
            'timestamp': datetime.now().isoformat(),
            'document_info': analysis.get('document_info', {}),
            'formatting_analysis': {
                'overall_score': analysis.get('compliance_score', 0),
                'section_breakdown': [
                    {
                        'name': section.name,
                        'formatting_quality': self._calculate_section_quality(section),
                        'issues': [
                            {
                                'severity': issue.severity,
                                'description': issue.description,
                                'suggestion': issue.suggestion
                            }
                            for issue in section.formatting_issues
                        ]
                    }
                    for section in analysis.get('sections', [])
                ],
                'style_consistency': analysis.get('style_consistency', {}),
                'visual_elements': analysis.get('visual_elements', [])
            },
            'recommendations': analysis.get('recommendations', []),
            'statistics': analysis.get('overall_statistics', {})
        }

    def _calculate_section_quality(self, section: SectionAnalysis) -> float:
        """Calculate quality score for individual section"""
        # Implementation of section quality scoring
        base_score = 1.0
        deductions = 0.0
        
        # Deduct for formatting issues
        for issue in section.formatting_issues:
            if issue.severity == 'high':
                deductions += 0.2
            elif issue.severity == 'medium':
                deductions += 0.1
            else:
                deductions += 0.05
                
        return max(0, min(1, base_score - deductions))

if __name__ == "__main__":
    # Example usage
    analyzer = ResumeFormatAnalyzer()
    
    # Analyze PDF resume
    pdf_analysis = analyzer.analyze_pdf_formatting("sample_resume.pdf")
    if pdf_analysis:
        report = analyzer.generate_report(pdf_analysis)
        print(json.dumps(report, indent=2, default=str))
    
    # Analyze DOCX resume
    docx_analysis = analyzer.analyze_docx_formatting("sample_resume.docx")
    if docx_analysis:
        report = analyzer.generate_report(docx_analysis)
        print(json.dumps(report, indent=2, default=str))
        
    def _calculate_font_score(self, analysis: Dict) -> float:
        """Calculate font consistency score"""
        try:
            fonts_used = analysis['overall_statistics'].get('unique_fonts', [])
            if not fonts_used:
                return 0.0
                
            # Check if fonts are from recommended list
            recommended_fonts = set(self.config['style_guidelines']['recommended_fonts'])
            recommended_font_usage = sum(1 for font in fonts_used if font in recommended_fonts)
            
            # Calculate scores
            font_count_score = min(1.0, self.config['formatting_rules']['max_fonts'] / len(fonts_used))
            font_choice_score = recommended_font_usage / len(fonts_used)
            
            return (font_count_score * 0.6 + font_choice_score * 0.4)
            
        except Exception as e:
            self.logger.error(f"Error calculating font score: {str(e)}")
            return 0.0

    def _calculate_spacing_score(self, analysis: Dict) -> float:
        """Calculate spacing consistency score"""
        try:
            avg_spacing = analysis['overall_statistics'].get('average_line_spacing', 0)
            if not avg_spacing:
                return 0.0
                
            recommended_spacing = self.config['formatting_rules']['recommended_line_spacing']
            
            # Calculate deviation from recommended spacing
            spacing_deviation = abs(avg_spacing - recommended_spacing) / recommended_spacing
            return max(0, 1 - spacing_deviation)
            
        except Exception as e:
            self.logger.error(f"Error calculating spacing score: {str(e)}")
            return 0.0

    def _calculate_section_score(self, analysis: Dict) -> float:
        """Calculate section structure score"""
        try:
            required_sections = set(self.config['section_requirements']['required'])
            found_sections = {section.name.lower() for section in analysis.get('sections', [])}
            
            # Calculate completeness
            section_completeness = len(required_sections.intersection(found_sections)) / len(required_sections)
            
            # Calculate order correctness
            section_order_score = self._check_section_order(analysis.get('sections', []))
            
            return (section_completeness * 0.7 + section_order_score * 0.3)
            
        except Exception as e:
            self.logger.error(f"Error calculating section score: {str(e)}")
            return 0.0

    def _calculate_visual_score(self, analysis: Dict) -> float:
        """Calculate visual appeal score"""
        try:
            scores = []
            
            # Check margins
            margin_score = self._check_margins(analysis)
            scores.append(margin_score)
            
            # Check alignment
            alignment_score = self._check_alignment(analysis)
            scores.append(alignment_score)
            
            # Check visual density
            density_score = self._check_visual_density(analysis)
            scores.append(density_score)
            
            return sum(scores) / len(scores) if scores else 0.0
            
        except Exception as e:
            self.logger.error(f"Error calculating visual score: {str(e)}")
            return 0.0

    def _check_margins(self, analysis: Dict) -> float:
        """Check if margins are within recommended range"""
        try:
            margins = analysis.get('document_info', {}).get('margins', {})
            if not margins:
                return 0.0
                
            min_margin = self.config['formatting_rules']['min_margin']
            max_margin = self.config['formatting_rules']['max_margin']
            
            margin_scores = []
            for margin in margins.values():
                if margin < min_margin:
                    score = margin / min_margin
                elif margin > max_margin:
                    score = max_margin / margin
                else:
                    score = 1.0
                margin_scores.append(score)
                
            return sum(margin_scores) / len(margin_scores)
            
        except Exception as e:
            self.logger.error(f"Error checking margins: {str(e)}")
            return 0.0

    def _check_alignment(self, analysis: Dict) -> float:
        """Check text alignment consistency"""
        try:
            alignment_counts = analysis['overall_statistics'].get('alignment_counts', {})
            if not alignment_counts:
                return 0.0
                
            # Calculate consistency score
            total_blocks = sum(alignment_counts.values())
            dominant_alignment = max(alignment_counts.values())
            
            return dominant_alignment / total_blocks
            
        except Exception as e:
            self.logger.error(f"Error checking alignment: {str(e)}")
            return 0.0

    def _check_visual_density(self, analysis: Dict) -> float:
        """Check if content density is appropriate"""
        try:
            text_area = analysis['overall_statistics'].get('text_area', 0)
            page_area = analysis['overall_statistics'].get('page_area', 0)
            
            if not page_area:
                return 0.0
                
            density = text_area / page_area
            
            # Ideal density range: 40-60%
            if 0.4 <= density <= 0.6:
                return 1.0
            elif density < 0.4:
                return density / 0.4
            else:
                return (1 - density) / 0.4
                
        except Exception as e:
            self.logger.error(f"Error checking visual density: {str(e)}")
            return 0.0

    def _check_section_order(self, sections: List[SectionAnalysis]) -> float:
        """Check if sections are in recommended order"""
        try:
            recommended_order = {
                section: idx 
                for idx, section in enumerate(self.config['section_requirements']['required'])
            }
            
            # Check order of found sections
            current_order = [
                section.name.lower()
                for section in sections
                if section.name.lower() in recommended_order
            ]
            
            if not current_order:
                return 0.0
                
            # Count inversions
            inversions = 0
            for i in range(len(current_order)):
                for j in range(i + 1, len(current_order)):
                    if recommended_order[current_order[i]] > recommended_order[current_order[j]]:
                        inversions += 1
                        
            max_inversions = (len(current_order) * (len(current_order) - 1)) / 2
            return 1 - (inversions / max_inversions if max_inversions else 0)
            
        except Exception as e:
            self.logger.error(f"Error checking section order: {str(e)}")
            return 0.0

    def detect_section_type(self, content: str) -> Optional[str]:
        """Detect type of section from content"""
        try:
            content_lower = content.lower()
            
            # Check against section patterns
            for pattern in self.section_patterns['headers']:
                if pattern.search(content_lower):
                    return pattern.pattern.strip(r'\b').lower()
                    
            # Additional heuristics
            if any(pattern.search(content) for pattern in self.section_patterns['contact_info']):
                return 'contact'
                
            return None
            
        except Exception as e:
            self.logger.error(f"Error detecting section type: {str(e)}")
            return None

    def analyze_image_block(self, block: Dict, page) -> Dict:
        """Analyze image block properties"""
        try:
            image_info = {
                'type': 'image',
                'position': (block.get('bbox', [0, 0, 0, 0])),
                'size': {
                    'width': block.get('width', 0),
                    'height': block.get('height', 0)
                }
            }
            
            # Extract image if possible
            try:
                image_data = page.get_images(block)
                if image_data:
                    image_info['format'] = image_data[0][1]
                    image_info['color_space'] = image_data[0][2]
            except:
                pass
                
            return image_info
            
        except Exception as e:
            self.logger.error(f"Error analyzing image block: {str(e)}")
            return {}

    def analyze_page_layout(self, page) -> List[FormattingIssue]:
        """Analyze page layout and identify issues"""
        try:
            issues = []
            
            # Check margins
            margins = self._get_page_margins(page)
            for side, margin in margins.items():
                if margin < self.config['formatting_rules']['min_margin']:
                    issues.append(FormattingIssue(
                        severity='medium',
                        category='margins',
                        description=f"{side.capitalize()} margin is too small",
                        suggestion=f"Increase {side} margin to at least {self.config['formatting_rules']['min_margin']} inches"
                    ))
            
            # Check content density
            density = self._calculate_content_density(page)
            if density > 0.8:
                issues.append(FormattingIssue(
                    severity='medium',
                    category='density',
                    description="Page content is too dense",
                    suggestion="Consider spreading content across multiple pages or reducing content"
                ))
            
            return issues
            
        except Exception as e:
            self.logger.error(f"Error analyzing page layout: {str(e)}")
            return []

    def analyze_style_consistency(self, doc) -> Dict:
        """Analyze style consistency across document"""
        try:
            style_analysis = {
                'paragraph_styles': defaultdict(int),
                'character_styles': defaultdict(int),
                'font_usage': defaultdict(int),
                'spacing_consistency': defaultdict(int)
            }
            
            for paragraph in doc.paragraphs:
                style_analysis['paragraph_styles'][paragraph.style.name] += 1
                
                # Analyze runs within paragraph
                for run in paragraph.runs:
                    if run.font.name:
                        style_analysis['font_usage'][run.font.name] += 1
                    if run.font.size:
                        style_analysis['character_styles'][f"size_{run.font.size}"] += 1
            
            return dict(style_analysis)
            
        except Exception as e:
            self.logger.error(f"Error analyzing style consistency: {str(e)}")
            return {}

    def _get_page_margins(self, page) -> Dict[str, float]:
        """Get page margins in inches"""
        try:
            # Convert page coordinates to inches
            rect = page.rect
            return {
                'left': rect.x0 / 72,  # Convert points to inches
                'right': (612 - rect.x1) / 72,  # Assuming standard page width
                'top': rect.y0 / 72,
                'bottom': (792 - rect.y1) / 72  # Assuming standard page height
            }
        except Exception as e:
            self.logger.error(f"Error getting page margins: {str(e)}")
            return {'left': 0, 'right': 0, 'top': 0, 'bottom': 0}

    def _calculate_content_density(self, page) -> float:
        """Calculate content density ratio"""
        try:
            text_blocks = page.get_text("dict")["blocks"]
            text_area = sum(
                (block['bbox'][2] - block['bbox'][0]) * (block['bbox'][3] - block['bbox'][1])
                for block in text_blocks
                if block.get('type') == 0
            )
            page_area = page.rect.width * page.rect.height
            return text_area / page_area if page_area else 0
            
        except Exception as e:
            self.logger.error(f"Error calculating content density: {str(e)}")
            return 0.0

    def get_pdf_document_info(self, doc) -> Dict:
        """Get PDF document metadata and properties"""
        try:
            return {
                'page_count': len(doc),
                'metadata': doc.metadata,
                'format': {
                    'width': doc[0].rect.width,
                    'height': doc[0].rect.height
                },
                'creation_date': doc.metadata.get('creationDate', ''),
                'modification_date': doc.metadata.get('modDate', '')
            }
        except Exception as e:
            self.logger.error(f"Error getting PDF document info: {str(e)}")
            return {}

    def get_docx_document_info(self, doc) -> Dict:
        """Get DOCX document metadata and properties"""
        try:
            core_properties = doc.core_properties
            return {
                'page_count': len(doc.sections),
                'metadata': {
                    'author': core_properties.author,
                    'created': core_properties.created,
                    'modified': core_properties.modified,
                    'title': core_properties.title
                },
                'sections': len(doc.sections),
                'paragraph_count': len(doc.paragraphs)
            }
        except Exception as e:
            self.logger.error(f"Error getting DOCX document info: {str(e)}")
            return {}